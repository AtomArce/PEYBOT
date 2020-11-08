from selenium import webdriver
import re
from docx import Document
from docx2pdf import convert


browser = webdriver.Chrome(executable_path= "d:\chromedriver.exe")

#======================== Open Portal ===============================#
browser.get("https://www.uoftengcareerportal.ca/students/login.htm")

def get_creds():

    id = input("Enter Student ID: \n")
    password = input("Enter password: \n")
    return id,password

# id, password = get_creds()
id, password = '1004476550','arceatom'



#=================== Login Page ========================#
student_id_ID = "j_username"
password_ID = "j_password"
login_button_xpath = '//*[@id="loginForm"]/div[3]/input'


#enter in credential keys#
browser.find_element_by_id(student_id_ID).send_keys(id)
browser.find_element_by_id(password_ID).send_keys(password)

browser.find_element_by_xpath(login_button_xpath).click()

# Click 'Job Postings' button
browser.find_element_by_xpath('//*[@id="mainContentDiv"]/div[2]/div/div[1]/div/a[2]').click();

# Click
browser.find_element_by_xpath('//*[@id="searchPostings"]/div[2]/div/ul/li/a').click();

# Click 'For My Program' button
browser.find_element_by_xpath('//*[@id="dashboard"]/div[2]/div[1]/a').click();


#========================================= filtering jobs===========================#

#==========================collect all job ids===========================#

#list of all job_ids
job_ids = []

col_1 = browser.find_elements_by_xpath('//*[@id="postingsTable"]/tbody//td[1]')  #selenium elements
for id in col_1:
    job_ids.append(id.text)  #ids are string type

#return to overview
browser.find_element_by_xpath('//*[@id="mainContentDiv"]/div[1]/div/div[2]/ul/li/a').click();

########################Cover Letter Stuff######################
def docx_replace_regex(doc_obj, regex , replace):

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex , replace)

def cover_letter_gen():
    regex1 = re.compile(r"job_name")
    regex2 = re.compile(r"company_name")
    replace1 = job_name
    replace2 = company_name
    doc_name = "coverletter.docx"
    doc = Document(doc_name)
    docx_replace_regex(doc, regex1, replace1)
    docx_replace_regex(doc, regex2, replace2)
    file_name = "CoverLetter-"+ id
    doc.save(file_name + ".docx")

    #Generates pdf file
    convert(file_name + ".docx")
    convert(file_name + ".docx", file_name + ".pdf")
#########################################################################

applied = []
external = []

#=========================== loop through job ids and apply if possible===========================#
for id in job_ids:
    #generate cover letter name & cover letter file

    # go to search bar for job id
    element = browser.find_element_by_xpath('// *[ @ id = "searchByPostingNumberForm"] / input[2]')
    browser.execute_script("arguments[0].scrollIntoView();", element)
    element.click();

    # #Click on apply button
    # browser.find_element_by_xpath('//*[@id="applyButton"]').click()

    # input id into bar
    browser.find_element_by_xpath('//*[@id="searchByPostingNumberForm"]/input[2]').send_keys(id)
    # click on search
    browser.find_element_by_xpath('//*[@id="searchByPostingNumberForm"]/a').click()

    job_name = browser.find_element_by_xpath('//*[@id="mainContentDiv"]/div[1]/div/div[2]/h1').text
    company_name = browser.find_element_by_xpath('//*[@id="postingDiv"]/table[3]/tbody/tr[1]/td[2]').text

    #click on apply for this position
    browser.find_element_by_xpath('//*[@id="applyButton"]').click()


    try: #check if I intend to apply exist->external
        browser.find_element_by_xpath('//*[@id="mainContentDiv"]/div[2]/div/div/div/div[2]/div/div/form/input[3]')
        external.append(id)  #add to external id list
        browser.find_element_by_xpath('//*[@id="mainContentDiv"]/div[1]/div/div[3]/ul/li/a').click()


    except: #apply case
        #click on upload document
        browser.find_element_by_xpath('//*[@id="mainContentDiv"]/div[2]/div/div/div/div[2]/div/div/div[2]/div[1]/a').click();
        #input cover letter name
        browser.find_element_by_xpath('//*[@id="fileUploadForm"]/div[1]/div[2]/input').send_keys(file_name)

        # click on document type
        browser.find_element_by_xpath('//*[@id="fileUploadForm"]/div[2]/div[2]/select').click()
        # click on cover letter option
        browser.find_element_by_xpath('//*[@id="fileUploadForm"]/div[2]/div[2]/select/option[2]').click()
        #click on choose file

        cover_letter_gen()

        browser.find_element_by_xpath('//*[@id="fileUploadForm"]/div[3]/div[2]/input').send_keys(r"D:\Users\Atom\PycharmProjects\PEYBOT\CoverLetter-" + id + ".pdf")
        #click on upload
        browser.find_element_by_xpath('//*[@id="mainContentDiv"]/div[2]/div/div/div/div/div/a[1]').click()

        break






